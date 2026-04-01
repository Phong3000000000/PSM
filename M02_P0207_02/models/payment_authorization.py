# -*- coding: utf-8 -*-
import io
import base64
from odoo import models, fields, api, _
from odoo.exceptions import UserError

class PaymentAuthorization(models.Model):
    _name = 'payment.authorization'
    _description = 'Payment Authorization (Ủy Nhiệm Chi)'
    _inherit = ['mail.thread', 'mail.activity.mixin']

    name = fields.Char(string='Reference', required=True, copy=False, readonly=True, default=lambda self: _('New'))
    date = fields.Date(string='Date', default=fields.Date.context_today, required=True, tracking=True)
    period = fields.Char(string='Period', required=True, tracking=True)
    
    total_amount = fields.Float(string='Total Amount', compute='_compute_total_amount', store=True, tracking=True)
    amount_text = fields.Char(string='Amount in Words', compute='_compute_amount_text', store=True)
    
    state = fields.Selection([
        ('draft', 'Draft'),
        ('pending', 'Pending Approval'),
        ('approved', 'Approved'),
        ('rejected', 'Rejected'),
        ('completed', 'Completed')
    ], string='Status', default='draft', tracking=True)

    payslip_ids = fields.One2many('hr.payslip', 'payment_order_id', string='Payslips')
    
    bank_id = fields.Many2one('res.partner.bank', string='Company Bank Account', domain="[('partner_id', '=', company_partner_id)]")
    company_id = fields.Many2one('res.company', string='Company', required=True, default=lambda self: self.env.company)
    company_partner_id = fields.Many2one('res.partner', related='company_id.partner_id', string='Company Partner')
    
    # Link to Payslip Run
    payslip_run_id = fields.Many2one('hr.payslip.run', string='Payslip Run', readonly=True)

    # Link to Approval Request (C-Level)
    approval_request_id = fields.Many2one(
        'approval.request', string='Approval Request (C-Level)',
        readonly=True, copy=False,
        help='Approval request gửi C-Level duyệt Payment Authorization này.'
    )

    # Link tới file UNC .docx đã tạo
    unc_attachment_id = fields.Many2one(
        'ir.attachment', string='File UNC (.docx)',
        readonly=True, copy=False,
    )

    def action_download_unc(self):
        """Smart button: tải file UNC .docx."""
        self.ensure_one()
        if not self.unc_attachment_id:
            return
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/content/%d?download=true' % self.unc_attachment_id.id,
            'target': 'self',
        }

    def action_view_approval_request(self):
        """Smart button: mở approval.request liên kết (C-Level)."""
        self.ensure_one()
        if not self.approval_request_id:
            return
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'approval.request',
            'view_mode': 'form',
            'res_id': self.approval_request_id.id,
            'target': 'current',
        }

    reason = fields.Char(string='Nội dung thanh toán', default='Thanh toán lương')

    # ── Section I: Phần đầu biểu mẫu ─────────────────────────────────────────
    x_ref_number    = fields.Char(string='Mã số hồ sơ')
    x_form_number   = fields.Char(string='Mẫu số', default='17')
    x_form_symbol   = fields.Char(string='Ký hiệu', default='C4-02aKB')
    x_voucher_number = fields.Char(string='Số chứng từ')
    x_qr_code       = fields.Char(string='Mã QR code')
    x_payment_method = fields.Selection([
        ('bank', 'Chuyển khoản'),
        ('cash', 'Tiền mặt'),
    ], string='Phương thức', default='bank')

    # ── Section II: Thông tin đơn vị trả tiền ────────────────────────────────
    x_payer_name          = fields.Char(string='Đơn vị trả tiền')
    x_payer_sponsor_code  = fields.Char(string='Mã nhà tài trợ')
    x_payer_address       = fields.Char(string='Địa chỉ đơn vị trả tiền')
    x_payer_account       = fields.Char(string='Tài khoản KBNN')
    x_payer_kbnn          = fields.Char(string='Tại KBNN')

    # ── Section III: Thông tin pháp lý ───────────────────────────────────────
    x_legal_decision_no     = fields.Char(string='QĐ phê duyệt số')
    x_legal_decision_day    = fields.Char(string='Ngày', size=3)
    x_legal_decision_month  = fields.Char(string='Tháng', size=3)
    x_legal_decision_year   = fields.Char(string='Năm', size=5)

    x_legal_contract_no     = fields.Char(string='HĐ/QĐ số')
    x_legal_contract_day    = fields.Char(string='Ngày', size=3)
    x_legal_contract_month  = fields.Char(string='Tháng', size=3)
    x_legal_contract_year   = fields.Char(string='Năm', size=5)
    x_legal_contract_amount = fields.Char(string='Tổng số tiền HĐ')
    x_legal_contract_party  = fields.Char(string='Ký với đơn vị/tổ chức')
    x_legal_advance_rule    = fields.Char(string='Quy định mức tạm ứng')

    x_legal_accept_no       = fields.Char(string='BB nghiệm thu số')
    x_legal_accept_day      = fields.Char(string='Ngày', size=3)
    x_legal_accept_month    = fields.Char(string='Tháng', size=3)
    x_legal_accept_year     = fields.Char(string='Năm', size=5)
    x_legal_accept_amount   = fields.Char(string='Tổng số tiền NT')
    x_legal_accept_party    = fields.Char(string='Ký với đơn vị (NT)')

    # ── Section VII: Thông tin đơn vị/cá nhân nhận tiền ─────────────────────
    x_recv_name       = fields.Char(string='Đơn vị/cá nhân nhận tiền')
    x_recv_tax_code   = fields.Char(string='Mã số thuế')
    x_recv_account    = fields.Char(string='Tài khoản nhận')
    x_recv_kbnn      = fields.Char(string='Tại KBNN/ngân hàng')
    x_recv_cash_name  = fields.Char(string='Người nhận tiền mặt')
    x_recv_cccd       = fields.Char(string='Số CCCD/Căn cước')
    x_recv_cccd_date  = fields.Char(string='Ngày cấp')
    x_recv_cccd_place = fields.Char(string='Nơi cấp')

    # ── Section VIII: Kho bạc Nhà nước ghi ───────────────────────────────────
    x_kbnn_debit1     = fields.Char(string='1. Nợ tài khoản')
    x_kbnn_credit1    = fields.Char(string='Có tài khoản (1)')
    x_kbnn_debit2     = fields.Char(string='2. Nợ tài khoản')
    x_kbnn_credit2    = fields.Char(string='Có tài khoản (2)')
    x_kbnn_area_code  = fields.Char(string='Mã địa bàn hành chính')
    x_kbnn_cash_bank  = fields.Char(string='NH/KBNN nơi nhận tiền mặt')

    # Legacy receiver fields (kept for compatibility)
    receiver_name         = fields.Char(string='Receiver Name')
    receiver_bank_account = fields.Char(string='Receiver Bank Account')
    receiver_bank_name    = fields.Char(string='Receiver Bank Name')

    @api.model_create_multi
    def create(self, vals_list):
        for vals in vals_list:
            if vals.get('name', _('New')) == _('New'):
                vals['name'] = self.env['ir.sequence'].next_by_code('payment.authorization') or _('New')
        return super(PaymentAuthorization, self).create(vals_list)

    @api.depends('payslip_ids.net_wage')
    def _compute_total_amount(self):
        for rec in self:
            rec.total_amount = sum(payslip.net_wage for payslip in rec.payslip_ids)

    @api.depends('total_amount', 'company_id.currency_id')
    def _compute_amount_text(self):
        for rec in self:
            rec.amount_text = rec.company_id.currency_id.amount_to_text(rec.total_amount)

    def action_submit(self):
        if not self.payslip_ids:
            raise UserError(_('You cannot submit a payment authorization without payslips.'))
        self.write({'state': 'pending'})

    
    # --- Step B8: C-Level Approval ---
    # Many2many to store who must approve? Or who *has* approved?
    # Spec: "Tất cả C-Level phải duyệt".
    # Implementation: 
    # 1. Define who are C-Levels (Group? List of Users?) -> Using 'res.groups' xmlref or config.
    # 2. Track approvals.
    
    approver_ids = fields.Many2many('res.users', 'payment_auth_approver_rel', 'payment_id', 'user_id', string='Approved By')
    required_approval_count = fields.Integer(compute='_compute_required_approvals', string='Required Approvals')
    approval_count = fields.Integer(compute='_compute_approval_count', string='Current Approvals')
    
    def _get_c_level_group(self):
        return self.env.ref('M02_P0207_02.group_c_level_approver', raise_if_not_found=False)

    @api.depends('state')
    def _compute_required_approvals(self):
        # Lấy từ cấu hình: số người phê duyệt C-Level cần thiết
        try:
            config = self.env['payroll.ops.config'].sudo().get_config()
            count = len(config.clevel_approver_ids)
        except Exception:
            count = 0
        for rec in self:
            rec.required_approval_count = max(count, 1)

    @api.depends('approver_ids')
    def _compute_approval_count(self):
        for rec in self:
            rec.approval_count = len(rec.approver_ids)

    def action_approve(self):
        """
        C-Level duyệt Payment Authorization.
        Khi đủ số người duyệt → validate payslips, gửi email, mark payslip run là completed.
        """
        self.ensure_one()
        # TODO: re-enable C-Level group check after testing
        # if not self.env.user.has_group('M02_P0207_02.group_c_level_approver'):
        #     raise UserError(_("Bạn không có quyền duyệt (cần quyền C-Level)."))

        # Add to approvers
        if self.env.user not in self.approver_ids:
            self.write({'approver_ids': [(4, self.env.user.id)]})
        
        # Check nếu đủ số người duyệt → hoàn tất
        if self.approval_count >= self.required_approval_count:
            self.write({'state': 'approved'})

            # 1. Force payslips sang 'done' (bypass state machine — Odoo 19 không có 'verify' state)
            non_done_slips = self.payslip_ids.filtered(lambda s: s.state not in ('done', 'cancel'))
            if non_done_slips:
                try:
                    non_done_slips.sudo().action_payslip_done()
                except Exception:
                    # Fallback: force done qua write()
                    non_done_slips.sudo().write({'state': 'done'})

            # 2. Mark payslip run là completed
            if self.payslip_run_id and self.payslip_run_id.state != 'completed':
                self.payslip_run_id.write({'state': 'completed'})
                self.payslip_run_id.message_post(
                    body=_('C-Level đã duyệt Payment Authorization %s. Quy trình lương hoàn tất.') % self.name
                )

            # 3. Tạo văn bản UNC (.docx) và lưu vào Attachments
            attachment = self._generate_unc_docx()

            # 4. Gửi thông báo Finance (kèm file)
            self._notify_finance(attachment)

            # 5. Gửi email phiếu lương từng nhân viên
            self.action_send_payslips()

    # ═════════════════════════════════════════════════════════════════════════
    #  TẠO VĂN BẢN UNC (.docx) – THEO NGHỊ ĐỊ NH 30/2020/NĐ-CP
    # ═════════════════════════════════════════════════════════════════════════

    def _generate_unc_docx(self):
        """
        Tạo file Word (.docx) ủy nhiệm chi theo đúng quy tắc văn bản hành chính
        Việt Nam (Nghị định 30/2020/NĐ-CP + Thông tư 01/2011/TT-BNV):
          – Font: Times New Roman 14pt
          – Lề trái 3.0cm, lề phải 2.0cm, trên/dưới 2.0cm
          – Giãn dòng 1.5 lines
        Lưu dưới dạng ir.attachment đính kèm vào bản ghi này.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            # Tự động cài python-docx nếu chưa có
            import subprocess, sys
            try:
                subprocess.check_call(
                    [sys.executable, '-m', 'pip', 'install', 'python-docx'],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
                from docx import Document
                from docx.shared import Pt, Cm
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
            except Exception as install_err:
                self.message_post(body=_(
                    '⚠️ Không thể tạo file Word: cài python-docx thất bại (%s). '
                    'Chạy thủ công: pip install python-docx') % install_err)
                return None

        doc = Document()

        # ── Cài đặt trang ─────────────────────────────────────────────────────
        section = doc.sections[0]
        section.page_height  = Cm(29.7)   # A4
        section.page_width   = Cm(21.0)
        section.left_margin  = Cm(3.0)    # NĐ 30/2020
        section.right_margin = Cm(2.0)
        section.top_margin   = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        # Helper: paragraph định dạng chuẩn
        def add_para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                     size=14, space_before=0, space_after=6):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after  = Pt(space_after)
            p.paragraph_format.line_spacing  = Pt(21)  # 1.5 dòng ~21pt với font 14
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            # Đầuăm: cần set cả w:cs để hiển dúng tiếng Việt
            run._element.rPr.get_or_add_rFonts().set(qn('w:cs'), 'Times New Roman')
            return p

        # Helper: row trong bảng
        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = align
            run = p.add_run(text or '')
            run.bold = bold
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)

        date_obj = self.date or fields.Date.today()
        company  = self.env.company

        # ── Tiêu đề ─────────────────────────────────────────────────────────
        h = add_para(company.name or 'TÊN ĐƠN VỊ', bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
        add_para('Mẫu số: %s / Ký hiệu: %s' % (
                     self.x_form_number or '17',
                     self.x_form_symbol  or 'C4-02aKB'),
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, space_after=2)
        add_para('ỦY NHIỆM CHI', bold=True, size=16,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
        add_para('Ngày %s tháng %s năm %s' % (
                     date_obj.day, date_obj.month, date_obj.year),
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=13, space_after=10)

        # ── Phần I ───────────────────────────────────────────────────────────
        add_para('I. PHẦN ĐẦU BIỂU MẪu', bold=True, space_before=8)
        add_para('- Số chứng từ: %s' % (self.x_voucher_number or '……………'))
        add_para('- Mã số hồ sơ: %s' % (self.x_ref_number or '………………………'))
        add_para('- Phương thức: %s' % (
            'Chuyển khoản' if self.x_payment_method == 'bank' else 'Tiền mặt'))

        # ── Phần II ──────────────────────────────────────────────────────────
        add_para('II. THÔNG TIN ĐƠN VỊ TRẢ TIỀN', bold=True, space_before=8)
        add_para('- Đơn vị trả tiền: %s'  % (self.x_payer_name or company.name or '…'))
        add_para('- Địa chỉ: %s'          % (self.x_payer_address or company.street or '…'))
        add_para('- Tài khoản KBNN: %s'   % (self.x_payer_account or '……………………………'))
        add_para('- Tại KBNN: %s'          % (self.x_payer_kbnn or '……………………………'))
        if self.bank_id:
            add_para('- Tài khoản ngân hàng công ty: %s (%s)' % (
                self.bank_id.acc_number, self.bank_id.bank_id.name or ''))

        # ── Phần IV: Nội dung ─────────────────────────────────────────────
        add_para('IV. NỘI DUNG THANH TOÁN', bold=True, space_before=8)
        add_para(self.reason or 'Thanh toán lương')

        # ── Phần V: Bảng chi tiết ──────────────────────────────────────────
        add_para('V. BẢNG CHI TIẾT THANH TOÁN', bold=True, space_before=8, space_after=4)

        headers = [
            'STT', 'Nhân viên', 'Người nhận',
            'Tài khoản', 'Tại ngân hàng',
            'Mã nguồn NS', 'Niên độ', 'Net Lương'
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            set_cell(hdr_cells[i], h, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

        # Data rows
        total = 0.0
        for idx, slip in enumerate(self.payslip_ids.filtered(lambda s: s.state != 'cancel'), 1):
            row_cells = table.add_row().cells
            net = slip.net_wage or 0.0
            total += net
            values = [
                str(idx),
                slip.employee_id.name or '',
                slip.x_receiver_name or '',
                slip.x_receiver_bank_account or '',
                slip.x_receiver_bank_name or '',
                slip.x_ma_nguon_ns or '',
                slip.x_nien_do_ns or str(date_obj.year),
                '{:,.0f}'.format(net),
            ]
            for i, val in enumerate(values):
                align = WD_ALIGN_PARAGRAPH.RIGHT if i == 7 else WD_ALIGN_PARAGRAPH.LEFT
                set_cell(row_cells[i], val, align=align, size=11)

        # Tổng cộng
        total_row = table.add_row().cells
        set_cell(total_row[0], 'Tổng cộng:', bold=True,
                 align=WD_ALIGN_PARAGRAPH.RIGHT, size=11)
        # Merge cells 0-6
        total_row[0].merge(total_row[6])
        set_cell(total_row[7], '{:,.0f}'.format(total), bold=True,
                 align=WD_ALIGN_PARAGRAPH.RIGHT, size=11)

        # ── Phần VI: Tổng tiền ───────────────────────────────────────────
        add_para('VI. TỔNG SỐ TIỀN BẰNG CHỮ:', bold=True, space_before=10)
        add_para(self.amount_text or '{:,.0f} đồng'.format(total),
                 italic=True)

        # ── Chữ ký ─────────────────────────────────────────────────────────
        add_para('IX – X. NGÀY XÁC NHẬN & CHỮ KÝ', bold=True, space_before=14)
        sig_table = doc.add_table(rows=3, cols=3)
        sig_table.style = 'Table Grid'
        roles = [['Kế toán trưởng', 'C-Level (Chủ tài khoản)', 'Xác nhận của KBNN'],
                 ['(Ký, ghi họ tên)',   '(Ký, ghi họ tên, đóng dấu)', '(Ký, họ tên, đóng dấu)'],
                 ['', '', '']]
        for r, row_texts in enumerate(roles):
            row_cells = sig_table.rows[r].cells
            for c, txt in enumerate(row_texts):
                set_cell(row_cells[c], txt,
                         bold=(r == 0),
                         align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

        # Gãn dòng cho chữ ký row
        row_cells = sig_table.rows[2].cells
        for c in range(3):
            row_cells[c].paragraphs[0].paragraph_format.space_before = Pt(40)

        # ── Lưu ra bytes ───────────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        docx_bytes = buf.read()

        filename = 'UNC_%s.docx' % (self.name or 'unc').replace('/', '_')
        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(docx_bytes).decode(),
            'res_model': 'payment.authorization',
            'res_id': self.id,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        })

        # Ghi lại link để smart button truy cập
        self.write({'unc_attachment_id': attachment.id})

        # Lưu vào Odoo Documents (module documents phải được cài)
        # Odoo 19: folder là documents.document với type='folder'
        if 'documents.document' in self.env:
            try:
                folder = self.env['documents.document'].sudo().search(
                    [('name', '=', 'Payroll'), ('type', '=', 'folder')], limit=1
                )
                if not folder:
                    # Tạo folder Payroll nếu chưa có
                    folder = self.env['documents.document'].sudo().create({
                        'name': 'Payroll',
                        'type': 'folder',
                        'access_internal': 'edit',
                    })
                self.env['documents.document'].sudo().create({
                    'name': filename,
                    'attachment_id': attachment.id,
                    'folder_id': folder.id,
                    'access_internal': 'edit',
                })
            except Exception as e:
                # Không chặn luồng chính nếu Documents integration lỗi
                import logging
                _logger = logging.getLogger(__name__)
                _logger.warning('Không lưu được vào Odoo Documents: %s', e)

        self.message_post(
            body=_('📄 File UNC đã tạo: <b>%s</b> — '
                   '<a href="/web/content/%d?download=true">⬇️ Tải về</a>') % (
                       filename, attachment.id),
            attachment_ids=[attachment.id],
        )
        return attachment


    def _notify_finance(self, attachment=None):
        """
        Gửi thông báo cho Finance khi PA được C-Level duyệt.
        – Luôn post chatter (notify followers của PA)
        – Nếu có finance_notify_ids trong config: tag họ vào message để gửi email
        – Kèm file UNC.docx nếu có
        """
        config = self.env['payroll.ops.config'].sudo().get_config()
        finance_users = config.finance_notify_ids
        finance_partners = finance_users.mapped('partner_id') if finance_users else self.env['res.partner']

        body = _(
            '<p>Xin chào,</p>'
            '<p>Payment Authorization <b>%(name)s</b> (Kỳ: %(period)s) '
            'đã được C-Level phê duyệt.</p>'
            '<p>Tổng số tiền: <b>%(amount)s</b></p>'
            '<p>Vui lòng xử lý chuyển khoản lương. 📄 File UNC đính kèm.</p>'
        ) % {
            'name':   self.name,
            'period': self.period or '',
            'amount': self.amount_text or '{:,.0f} VNĐ'.format(self.total_amount),
        }

        self.message_post(
            body=body,
            partner_ids=finance_partners.ids,
            subject=_('[Lương] UNC %s đã duyệt - cần xử lý chuyển khoản') % self.name,
            message_type='comment',
            subtype_xmlid='mail.mt_comment',
            attachment_ids=[attachment.id] if attachment else [],
        )

    def action_reject_line(self, line_id, reason):
        """ Loopback: Reject specific line """
        pass

    def action_send_payslips(self):
        """
        Gửi email phiếu lương cho từng nhân viên.
        Dùng template chuẩn hr_payroll.mail_template_new_payslip.
        Chỉ gửi cho payslip ở state 'done' và employee có work_email.
        """
        template = self.env.ref('hr_payroll.mail_template_new_payslip', raise_if_not_found=False)
        if not template:
            self.message_post(body=_('Không tìm thấy email template phiếu lương. Bỏ qua bước gửi email nhân viên.'))
            return

        sent_count = 0
        failed_count = 0
        for slip in self.payslip_ids.filtered(lambda s: s.state != 'cancel'):
            email = slip.employee_id.work_email
            if not email:
                failed_count += 1
                continue
            try:
                template.send_mail(
                    slip.id,
                    email_values={'email_to': email},
                    force_send=True,
                )
                sent_count += 1
            except Exception as e:
                failed_count += 1
                # Log but don't block
                slip.message_post(body=_('Không gửi được email: %s') % str(e))

        self.message_post(
            body=_('Đã gửi phiếu lương qua email: %(sent)d thành công, %(failed)d thất bại.') % {
                'sent': sent_count,
                'failed': failed_count,
            }
        )


    def action_reject(self):
        self.write({'state': 'rejected'})
        # Reset linked payslips? Logic defined in requirement 8.2 (Workflow Rollback)
        # "Khi quay lại, tất cả payslip và payment_authorization chuyển sang state='rejected'"
        # self.payslip_ids.write({'state': 'rejected'})

    def action_complete(self):
        self.write({'state': 'completed'})
        # Update payslip state to paid
        # self.payslip_ids.write({'state': 'paid'})
    
    def action_draft(self):
        self.write({'state': 'draft'})

